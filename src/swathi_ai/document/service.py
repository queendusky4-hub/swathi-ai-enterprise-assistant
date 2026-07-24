from __future__ import annotations

import io
import math
import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from threading import RLock
from uuid import uuid4


# ---------------------------------------------------------------------
# DATA MODELS
# ---------------------------------------------------------------------


@dataclass(slots=True)
class DocumentChunk:
    chunk_id: str
    document_id: str
    filename: str
    text: str
    page_number: int | None = None
    section_type: str | None = None
    chunk_index: int = 0


@dataclass(slots=True)
class DocumentIndexResult:
    document_id: str
    filename: str
    chunk_count: int


@dataclass(slots=True)
class DocumentSearchResult:
    chunk_id: str
    document_id: str
    filename: str
    text: str
    score: float
    page_number: int | None = None
    section_type: str | None = None
    chunk_index: int | None = None


# ---------------------------------------------------------------------
# DOCUMENT RAG SERVICE
# ---------------------------------------------------------------------


class DocumentRAGService:
    """
    Thread-safe, in-memory document indexing and retrieval service.

    Supported formats:
    - PDF
    - DOCX
    - TXT
    - Markdown
    - CSV

    Notes:
    - Documents are stored in memory.
    - Restarting FastAPI clears the index.
    - Use one shared service instance through get_document_service().
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 80,
    ) -> None:
        if chunk_size <= 0:
            raise ValueError("chunk_size must be greater than zero.")

        if chunk_overlap < 0:
            raise ValueError("chunk_overlap cannot be negative.")

        if chunk_overlap >= chunk_size:
            raise ValueError(
                "chunk_overlap must be smaller than chunk_size."
            )

        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        self._chunks: list[DocumentChunk] = []
        self._documents: dict[str, dict[str, object]] = {}
        self._lock = RLock()

    # -----------------------------------------------------------------
    # DOCUMENT INDEXING
    # -----------------------------------------------------------------

    def index_document(
        self,
        *,
        filename: str,
        file_bytes: bytes,
    ) -> DocumentIndexResult:
        """
        Extract, chunk, and index a document.
        """

        clean_filename = Path(filename or "").name.strip()

        if not clean_filename:
            raise ValueError("A valid filename is required.")

        if not file_bytes:
            raise ValueError("The uploaded document is empty.")

        extension = Path(clean_filename).suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            supported = ", ".join(
                sorted(self.SUPPORTED_EXTENSIONS)
            )
            raise ValueError(
                f"Unsupported document format '{extension}'. "
                f"Supported formats: {supported}."
            )

        extracted_pages = self._extract_document(
            filename=clean_filename,
            file_bytes=file_bytes,
        )

        if not extracted_pages:
            raise ValueError(
                "No readable content could be extracted from the document."
            )

        document_id = str(uuid4())
        created_chunks: list[DocumentChunk] = []
        chunk_index = 0

        for page_number, page_text in extracted_pages:
            clean_text = self._clean_text(page_text)

            if not clean_text:
                continue

            for chunk_text in self._split_text(clean_text):
                clean_chunk = chunk_text.strip()

                if not clean_chunk:
                    continue

                created_chunks.append(
                    DocumentChunk(
                        chunk_id=str(uuid4()),
                        document_id=document_id,
                        filename=clean_filename,
                        text=clean_chunk,
                        page_number=page_number,
                        section_type=self._detect_section_type(
                            clean_chunk
                        ),
                        chunk_index=chunk_index,
                    )
                )
                chunk_index += 1

        if not created_chunks:
            raise ValueError(
                "The document was read, but no usable chunks were created."
            )

        document_metadata: dict[str, object] = {
            "document_id": document_id,
            "filename": clean_filename,
            "chunk_count": len(created_chunks),
        }

        with self._lock:
            self._chunks.extend(created_chunks)
            self._documents[document_id] = document_metadata

        return DocumentIndexResult(
            document_id=document_id,
            filename=clean_filename,
            chunk_count=len(created_chunks),
        )

    # -----------------------------------------------------------------
    # DOCUMENT SEARCH
    # -----------------------------------------------------------------

    def search(
        self,
        query: str,
        top_k: int = 5,
        document_ids: list[str] | None = None,
    ) -> list[DocumentSearchResult]:
        """
        Search indexed document chunks.

        Retrieval uses:
        - cosine similarity
        - query token coverage
        - phrase matching
        - section-aware boosts
        - synonym expansion

        Matching document chunks are not discarded merely because their
        lexical score is zero. A very small fallback score keeps selected
        document context available for the answer generator.
        """

        clean_query = (query or "").strip()

        if not clean_query:
            return []

        if top_k <= 0:
            return []

        selected_document_ids = (
            {item.strip() for item in document_ids if item.strip()}
            if document_ids
            else None
        )

        with self._lock:
            chunks_snapshot = list(self._chunks)

        if not chunks_snapshot:
            return []

        expanded_query_tokens = self._expand_query_tokens(
            self._tokenize(clean_query)
        )
        query_vector = Counter(expanded_query_tokens)

        scored_results: list[DocumentSearchResult] = []

        for chunk in chunks_snapshot:
            if (
                selected_document_ids is not None
                and chunk.document_id not in selected_document_ids
            ):
                continue

            chunk_tokens = self._tokenize(chunk.text)
            chunk_vector = Counter(chunk_tokens)

            cosine_score = self._cosine_similarity(
                query_vector,
                chunk_vector,
            )
            coverage_score = self._query_coverage_score(
                expanded_query_tokens,
                chunk_tokens,
            )
            phrase_score = self._phrase_match_score(
                clean_query,
                chunk.text,
            )
            section_boost = self._section_boost(
                query=clean_query,
                section_type=chunk.section_type,
                text=chunk.text,
            )

            score = (
                cosine_score * 0.45
                + coverage_score * 0.35
                + phrase_score * 0.15
                + section_boost * 0.05
            )

            if score <= 0:
                score = 0.000001

            scored_results.append(
                DocumentSearchResult(
                    chunk_id=chunk.chunk_id,
                    document_id=chunk.document_id,
                    filename=chunk.filename,
                    text=chunk.text,
                    score=round(float(score), 6),
                    page_number=chunk.page_number,
                    section_type=chunk.section_type,
                    chunk_index=chunk.chunk_index,
                )
            )

        scored_results.sort(
            key=lambda item: (
                item.score,
                -(item.chunk_index or 0),
            ),
            reverse=True,
        )

        return scored_results[:top_k]

    # -----------------------------------------------------------------
    # DOCUMENT MANAGEMENT
    # -----------------------------------------------------------------

    def list_documents(self) -> list[dict[str, object]]:
        with self._lock:
            return [
                dict(item)
                for item in self._documents.values()
            ]

    def get_document(
        self,
        document_id: str,
    ) -> dict[str, object] | None:
        with self._lock:
            item = self._documents.get(document_id)
            return dict(item) if item else None

    def delete_document(self, document_id: str) -> bool:
        clean_document_id = (document_id or "").strip()

        if not clean_document_id:
            return False

        with self._lock:
            if clean_document_id not in self._documents:
                return False

            self._chunks = [
                chunk
                for chunk in self._chunks
                if chunk.document_id != clean_document_id
            ]
            del self._documents[clean_document_id]

        return True

    def clear(self) -> None:
        with self._lock:
            self._chunks.clear()
            self._documents.clear()

    def get_stats(self) -> dict[str, int]:
        with self._lock:
            return {
                "document_count": len(self._documents),
                "chunk_count": len(self._chunks),
            }

    # -----------------------------------------------------------------
    # FILE EXTRACTION
    # -----------------------------------------------------------------

    def _extract_document(
        self,
        *,
        filename: str,
        file_bytes: bytes,
    ) -> list[tuple[int | None, str]]:
        extension = Path(filename).suffix.lower()

        if extension == ".pdf":
            return self._extract_pdf(file_bytes)

        if extension == ".docx":
            return self._extract_docx(file_bytes)

        if extension in {".txt", ".md", ".csv"}:
            return self._extract_plain_text(file_bytes)

        raise ValueError(
            f"Unsupported document extension: {extension}"
        )

    @staticmethod
    def _extract_pdf(
        file_bytes: bytes,
    ) -> list[tuple[int | None, str]]:
        try:
            import fitz
        except ImportError as error:
            raise RuntimeError(
                "PDF support requires PyMuPDF. "
                "Install it with: pip install pymupdf"
            ) from error

        pages: list[tuple[int | None, str]] = []

        try:
            with fitz.open(
                stream=file_bytes,
                filetype="pdf",
            ) as document:
                for page_index, page in enumerate(document):
                    text = page.get_text("text") or ""

                    if text.strip():
                        pages.append(
                            (page_index + 1, text)
                        )
        except Exception as error:
            raise ValueError(
                f"Unable to read the PDF document: {error}"
            ) from error

        return pages

    @staticmethod
    def _extract_docx(
        file_bytes: bytes,
    ) -> list[tuple[int | None, str]]:
        try:
            from docx import Document
        except ImportError as error:
            raise RuntimeError(
                "DOCX support requires python-docx. "
                "Install it with: pip install python-docx"
            ) from error

        try:
            document = Document(io.BytesIO(file_bytes))
        except Exception as error:
            raise ValueError(
                f"Unable to read the DOCX document: {error}"
            ) from error

        content: list[str] = []

        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if value:
                content.append(value)

        for table in document.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    content.append(" | ".join(cells))

        text = "\n".join(content).strip()

        return [(None, text)] if text else []

    @staticmethod
    def _extract_plain_text(
        file_bytes: bytes,
    ) -> list[tuple[int | None, str]]:
        encodings = (
            "utf-8-sig",
            "utf-8",
            "utf-16",
            "cp1252",
            "latin-1",
        )

        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)

                if text.strip():
                    return [(None, text)]
            except UnicodeDecodeError:
                continue

        raise ValueError(
            "Unable to decode the uploaded text document."
        )

    # -----------------------------------------------------------------
    # CHUNKING
    # -----------------------------------------------------------------

    def _split_text(self, text: str) -> list[str]:
        """
        Split text into overlapping word-based chunks.
        """

        words = text.split()

        if not words:
            return []

        if len(words) <= self.chunk_size:
            return [" ".join(words)]

        chunks: list[str] = []
        start = 0

        while start < len(words):
            end = min(
                start + self.chunk_size,
                len(words),
            )

            chunk_text = " ".join(words[start:end]).strip()

            if chunk_text:
                chunks.append(chunk_text)

            if end >= len(words):
                break

            start = max(
                end - self.chunk_overlap,
                start + 1,
            )

        return chunks

    # -----------------------------------------------------------------
    # TEXT PROCESSING
    # -----------------------------------------------------------------

    @staticmethod
    def _clean_text(text: str) -> str:
        value = (text or "").replace("\x00", " ")
        value = value.replace("\r\n", "\n")
        value = value.replace("\r", "\n")
        value = re.sub(r"[ \t]+", " ", value)
        value = re.sub(r"\n{3,}", "\n\n", value)
        return value.strip()

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        return re.findall(
            r"[A-Za-z0-9_+#.\-]+|[\u0B80-\u0BFF]+",
            (text or "").lower(),
        )

    @staticmethod
    def _expand_query_tokens(
        tokens: list[str],
    ) -> list[str]:
        synonym_map = {
            "degree": {
                "degree",
                "qualification",
                "education",
                "msc",
                "master",
                "masters",
                "bsc",
                "bachelor",
                "bachelors",
                "university",
            },
            "qualification": {
                "qualification",
                "degree",
                "education",
                "msc",
                "master",
                "bsc",
                "bachelor",
            },
            "skills": {
                "skills",
                "technologies",
                "tools",
                "languages",
                "frameworks",
            },
            "experience": {
                "experience",
                "employment",
                "work",
                "role",
                "company",
            },
            "projects": {
                "projects",
                "portfolio",
                "application",
                "system",
            },
            "certifications": {
                "certifications",
                "certificate",
                "course",
                "training",
            },
        }

        expanded = list(tokens)

        for token in tokens:
            expanded.extend(
                synonym_map.get(token, set())
            )

        return expanded

    # -----------------------------------------------------------------
    # SCORING
    # -----------------------------------------------------------------

    @staticmethod
    def _cosine_similarity(
        first_vector: Counter[str],
        second_vector: Counter[str],
    ) -> float:
        if not first_vector or not second_vector:
            return 0.0

        common_tokens = (
            set(first_vector)
            & set(second_vector)
        )

        dot_product = sum(
            first_vector[token] * second_vector[token]
            for token in common_tokens
        )

        first_magnitude = math.sqrt(
            sum(
                value * value
                for value in first_vector.values()
            )
        )
        second_magnitude = math.sqrt(
            sum(
                value * value
                for value in second_vector.values()
            )
        )

        if first_magnitude == 0 or second_magnitude == 0:
            return 0.0

        return dot_product / (
            first_magnitude * second_magnitude
        )

    @staticmethod
    def _query_coverage_score(
        query_tokens: list[str],
        chunk_tokens: list[str],
    ) -> float:
        if not query_tokens or not chunk_tokens:
            return 0.0

        query_set = set(query_tokens)
        chunk_set = set(chunk_tokens)
        overlap = query_set.intersection(chunk_set)

        return (
            len(overlap) / len(query_set)
            if overlap
            else 0.0
        )

    @staticmethod
    def _phrase_match_score(
        query: str,
        text: str,
    ) -> float:
        clean_query = " ".join(
            (query or "").lower().split()
        )
        clean_text = " ".join(
            (text or "").lower().split()
        )

        if not clean_query or not clean_text:
            return 0.0

        if clean_query in clean_text:
            return 1.0

        return 0.0

    @staticmethod
    def _section_boost(
        *,
        query: str,
        section_type: str | None,
        text: str,
    ) -> float:
        lowered_query = (query or "").lower()
        lowered_text = (text or "").lower()

        section_terms = {
            "education": (
                "degree",
                "education",
                "qualification",
                "university",
                "msc",
                "master",
                "bsc",
                "bachelor",
            ),
            "experience": (
                "experience",
                "employment",
                "worked",
                "company",
                "role",
            ),
            "skills": (
                "skill",
                "skills",
                "technology",
                "technologies",
                "tools",
            ),
            "projects": (
                "project",
                "projects",
                "portfolio",
            ),
            "certifications": (
                "certification",
                "certificate",
                "course",
            ),
        }

        detected_section = section_type

        if not detected_section:
            detected_section = DocumentRAGService._detect_section_type(
                lowered_text
            )

        if not detected_section:
            return 0.0

        terms = section_terms.get(
            detected_section,
            (),
        )

        return (
            1.0
            if any(term in lowered_query for term in terms)
            else 0.0
        )

    @staticmethod
    def _detect_section_type(
        text: str,
    ) -> str | None:
        lowered = (text or "").lower()

        section_keywords = {
            "education": (
                "education",
                "academic",
                "university",
                "degree",
                "msc",
                "master",
                "bsc",
                "bachelor",
                "qualification",
            ),
            "experience": (
                "experience",
                "employment",
                "work history",
                "professional experience",
            ),
            "skills": (
                "skills",
                "technologies",
                "technical skills",
                "programming languages",
                "tools",
            ),
            "projects": (
                "projects",
                "project experience",
                "portfolio",
            ),
            "certifications": (
                "certification",
                "certifications",
                "certificate",
                "courses",
            ),
            "profile": (
                "profile",
                "summary",
                "objective",
                "about me",
            ),
        }

        for section_type, keywords in section_keywords.items():
            if any(
                keyword in lowered
                for keyword in keywords
            ):
                return section_type

        return None