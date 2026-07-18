from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}


class DocumentProcessingError(Exception):
    """Raised when an uploaded document cannot be processed."""


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract readable text from a PDF file."""

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as error:
        raise DocumentProcessingError(
            "The PDF file could not be opened."
        ) from error

    pages: list[str] = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""

        page_text = page_text.strip()

        if page_text:
            pages.append(
                f"[Page {page_number}]\n{page_text}"
            )

    if not pages:
        raise DocumentProcessingError(
            "No readable text was found in the PDF. "
            "The PDF may contain only scanned images."
        )

    return "\n\n".join(pages)


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract paragraphs and tables from a DOCX file."""

    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as error:
        raise DocumentProcessingError(
            "The DOCX file could not be opened."
        ) from error

    sections: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            sections.append(paragraph_text)

    for table_number, table in enumerate(
        document.tables,
        start=1,
    ):
        rows: list[str] = []

        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            sections.append(
                f"[Table {table_number}]\n"
                + "\n".join(rows)
            )

    if not sections:
        raise DocumentProcessingError(
            "No readable text was found in the DOCX file."
        )

    return "\n\n".join(sections)


def extract_txt_text(file_bytes: bytes) -> str:
    """Decode a plain-text file."""

    encodings = (
        "utf-8",
        "utf-8-sig",
        "utf-16",
        "cp1252",
    )

    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding).strip()

            if text:
                return text
        except UnicodeDecodeError:
            continue

    raise DocumentProcessingError(
        "The TXT file encoding could not be recognised."
    )


def extract_document_text(
    filename: str,
    file_bytes: bytes,
) -> str:
    """Extract text according to the uploaded file extension."""

    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        supported_formats = ", ".join(
            sorted(SUPPORTED_EXTENSIONS)
        )

        raise DocumentProcessingError(
            "Unsupported document type. "
            f"Supported formats: {supported_formats}"
        )

    if not file_bytes:
        raise DocumentProcessingError(
            "The uploaded document is empty."
        )

    if extension == ".pdf":
        return extract_pdf_text(file_bytes)

    if extension == ".docx":
        return extract_docx_text(file_bytes)

    if extension == ".txt":
        return extract_txt_text(file_bytes)

    raise DocumentProcessingError(
        "The document type could not be processed."
    )